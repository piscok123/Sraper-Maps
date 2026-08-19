# modules/plugins/review_scraper.py
import asyncio
import csv
import json
import os
import random
import re
from datetime import datetime
from typing import List, Dict, Any, Optional, Callable

from playwright.async_api import Page, BrowserContext, async_playwright, Browser
from rich.console import Console
from rich.progress import Progress, SpinnerColumn, TextColumn, BarColumn, TaskProgressColumn
import urllib.parse

import docx
from openpyxl import Workbook
from openpyxl.styles import Font
from openpyxl.utils import get_column_letter

console = Console()

USER_AGENTS = [
    "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36",
    "Mozilla/5.0 (Macintosh; Intel Mac OS X 10_15_7) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/119.0.0.0 Safari/537.36"
]

def add_hyperlink(paragraph, url, text):
    part = paragraph.part
    r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)
    hyperlink = docx.oxml.shared.OxmlElement('w:hyperlink')
    hyperlink.set(docx.oxml.shared.qn('r:id'), r_id, )
    new_run = docx.oxml.shared.OxmlElement('w:r')
    rPr = docx.oxml.shared.OxmlElement('w:rPr')
    c = docx.oxml.shared.OxmlElement('w:color')
    c.set(docx.oxml.shared.qn('w:val'), '0563C1')
    u = docx.oxml.shared.OxmlElement('w:u')
    u.set(docx.oxml.shared.qn('w:val'), 'single')
    rPr.append(c)
    rPr.append(u)
    new_run.append(rPr)
    new_run_text = docx.oxml.shared.OxmlElement('w:t')
    new_run_text.text = text
    new_run.append(new_run_text)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)
    return hyperlink


class GoogleMapsReviewScraperAsync:
    def __init__(self, target_url: str, limit: int, headless: bool, proxy: Optional[str] = None, output_dir: str = "output", browser: Optional[Browser] = None, progress_callback: Optional[Callable[[int, int, str], Any]] = None) -> None:
        self.target_url = target_url
        self.limit = limit
        self.headless = headless
        self.proxy_server = {"server": proxy} if proxy else None
        self.browser = browser
        self.results: List[Dict[str, Any]] = []
        self.output_dir = output_dir
        self.session_dir = os.path.join(os.getcwd(), "browser_session")
        self.place_name = "Google_Maps"
        self.progress_callback = progress_callback
        
        if not os.path.exists(self.output_dir): os.makedirs(self.output_dir)

        self.headers = ["Reviewer Name", "Rating", "Date/Time", "Review Text"]

    async def _report_progress(self, current: int, total: int, phase_text: str = "") -> None:
        if self.progress_callback:
            try:
                if asyncio.iscoroutinefunction(self.progress_callback):
                    await self.progress_callback(current, total, phase_text)
                else:
                    self.progress_callback(current, total, phase_text)
            except Exception:
                pass

    async def _human_delay(self, min_sec=1.0, max_sec=2.5): 
        await asyncio.sleep(random.uniform(min_sec, max_sec))

    async def scrape(self) -> None:
        if self.browser:
            context = await self.browser.new_context(
                viewport={"width": 1280, "height": 800},
                user_agent=USER_AGENTS[0],
                locale="id-ID"
            )
            await self._run_scrape(context)
            await context.close()
        else:
            async with async_playwright() as p:
                console.print("[bold cyan]🤖 Menyiapkan Browser Profile (Persistent Context)...[/bold cyan]")
                
                context = await p.chromium.launch_persistent_context(
                    user_data_dir=self.session_dir,
                    headless=self.headless,
                    args=['--disable-blink-features=AutomationControlled'],
                    proxy=self.proxy_server,
                    viewport={"width": 1280, "height": 800},
                    user_agent=USER_AGENTS[0],
                    locale="id-ID"
                )
                await self._run_scrape(context)
                await context.close()

    async def _run_scrape(self, context) -> None:
        async def _block_heavy_resources(route):
            if route.request.resource_type in ["image", "font", "media"]:
                await route.abort()
            else:
                await route.continue_()

        try:
            await context.route("**/*", _block_heavy_resources)
        except Exception:
            pass

        page = context.pages[0] if len(context.pages) > 0 else await context.new_page()

        console.print("[bold green]🔍 Membuka Google Maps...[/bold green]")
        try:
            clean_url = self.target_url.replace('\n', '').replace('\r', '').replace(' ', '').strip()
            if not clean_url.startswith('http'):
                clean_url = 'https://' + clean_url

            await page.goto(clean_url, wait_until="domcontentloaded", timeout=60000)
            await self._human_delay(3.0, 5.0)

            try:
                title_elem = page.locator('h1')
                if await title_elem.count() > 0:
                    self.place_name = (await title_elem.first.inner_text()).strip()
            except: pass

            try:
                cookie_btn = page.locator('button:has-text("Terima semua"), button:has-text("Accept all")')
                if await cookie_btn.count() > 0: await cookie_btn.first.click(timeout=3000)
            except: pass

            console.print("[cyan]Mengkondisikan Tab Ulasan...[/cyan]")
            
            # Coba klik tab ulasan hanya jika URL tidak langsung menuju tab ulasan
            if "!9m1!1b1" not in clean_url:
                await page.evaluate("""() => {
                    let tabs = Array.from(document.querySelectorAll('button[role="tab"]'));
                    let reviewTab = tabs.find(tab => tab.innerText.toLowerCase().includes('ulasan') || tab.innerText.toLowerCase().includes('reviews'));
                    if (reviewTab) {
                        reviewTab.click();
                    } else {
                        let btns = Array.from(document.querySelectorAll('button'));
                        let moreBtn = btns.find(b => b.innerText.toLowerCase().includes('ulasan lainnya') || b.innerText.toLowerCase().includes('more reviews'));
                        if (moreBtn) moreBtn.click();
                    }
                }""")
                await self._human_delay(3.0, 5.0)

            # Pastikan minimal 1 elemen review terlihat
            try:
                await page.locator('div[data-review-id]').first.wait_for(state="visible", timeout=15000)
            except:
                console.print("[bold red]Gagal menemukan blok Ulasan. Google mungkin merender UI yang berbeda.[/bold red]")
                return

            seen_reviews = set()
            retries = 0 

            console.print(f"\n[bold magenta]⚡ Memulai Ekstraksi (Membaca On-The-Fly sambil Scrolling)...[/bold magenta]")
            
            is_bot = os.environ.get('IS_TELEGRAM_BOT') == '1'
            if is_bot:
                target_limit = self.limit
                class DummyProgress:
                    def add_task(self, *args, **kwargs): return None
                    def update(self, *args, **kwargs):
                        completed = kwargs.get('completed', 0)
                        console.print(f"[cyan]Mengekstrak Ulasan... Ditemukan: {completed}/{target_limit}[/cyan]", end="\r")
                    def __enter__(self): return self
                    def __exit__(self, *args): console.print()
                progress_ctx = DummyProgress()
            else:
                progress_ctx = Progress(SpinnerColumn(), TextColumn("[cyan]Mengekstrak Ulasan... Ditemukan: {task.completed}/{task.total}"), BarColumn(), TaskProgressColumn())
                
            with progress_ctx as progress:
                task = progress.add_task("Mengekstrak", total=self.limit)
                
                while len(self.results) < self.limit:
                    # Dapatkan jumlah block saat ini di layar
                    blocks_count = await page.locator('div[data-review-id]').count()
                    
                    if blocks_count == 0:
                        break

                    extracted_in_this_cycle = 0

                    # --- JAVASCRIPT INJECTION: Ekstrak SEMUA Data Sekaligus ---
                    js_extract_code = """
                    (seenIds) => {
                        let results = [];
                        let blocks = document.querySelectorAll('div[data-review-id]');
                        for (let el of blocks) {
                            let id = el.getAttribute('data-review-id');
                            if (seenIds.includes(id)) continue;
                            
                            // Klik tombol "Lengkapnya" jika ada
                            let btn = Array.from(el.querySelectorAll('button')).find(b => b.innerText.toLowerCase().includes('lengkapnya') || b.innerText.toLowerCase().includes('more'));
                            if (btn) btn.click();
                            
                            let name = "";
                            let avatar = el.querySelector('button[aria-label]');
                            if (avatar) {
                                name = avatar.getAttribute('aria-label');
                            } else {
                                let btnFirst = el.querySelector('button');
                                if (btnFirst) name = btnFirst.innerText;
                            }
                            
                            let rating = "";
                            let ratingElem = el.querySelector('span[role="img"][aria-label*="bintang"], span[role="img"][aria-label*="stars"]');
                            if (ratingElem) rating = ratingElem.getAttribute('aria-label');
                            
                            let date = "";
                            let dateElem = el.querySelector('span.rsqaWe') || Array.from(el.querySelectorAll('span')).find(s => s.innerText.includes('lalu') || s.innerText.includes('ago'));
                            if (dateElem) date = dateElem.innerText;
                            
                            let text = "";
                            let textElem = el.querySelector('span.wiI7pd');
                            if (textElem) text = textElem.innerText;
                            
                            results.push({id, name, rating, date, text});
                        }
                        return results;
                    }
                    """
                    
                    extracted_list = await page.evaluate(js_extract_code, list(seen_reviews))
                    
                    for extracted in extracted_list:
                        if len(self.results) >= self.limit:
                            break
                            
                        review_id = extracted["id"]
                        if not review_id or review_id in seen_reviews:
                            continue
                            
                        seen_reviews.add(review_id)
                        
                        data = {k: "N/A" for k in self.headers}
                        
                        if extracted["name"]:
                            data["Reviewer Name"] = re.sub(r'^(Foto|Photo of)\s+', '', extracted["name"].strip(), flags=re.IGNORECASE)
                            
                        if extracted["rating"]:
                            match = re.search(r'(\d)', extracted["rating"])
                            if match: data["Rating"] = f"{match.group(1)} Bintang"
                            
                        if extracted["date"]:
                            data["Date/Time"] = extracted["date"].strip()
                            
                        if extracted["text"]:
                            data["Review Text"] = extracted["text"].strip()
                        
                        self.results.append(data)
                        extracted_in_this_cycle += 1
                        progress.update(task, completed=len(self.results))
                        await self._report_progress(len(self.results), self.limit, f"Mengekstrak ulasan ({len(self.results)}/{self.limit})...")

                    # FIX SCROLLING AGGRESIF: Menggunakan simulasi Hardware Asli (Mouse & Keyboard)
                    if blocks_count > 0:
                        try:
                            # 1. Cari elemen terakhir berdasarkan index (blocks_count - 1)
                            last_block = page.locator('div[data-review-id]').nth(blocks_count - 1)
                            await last_block.hover(timeout=2000)
                            
                            # 2. Gulir roda mouse ke bawah secara drastis
                            await page.mouse.wheel(0, 10000)
                            
                            # 3. Tekan tombol PageDown berulang kali untuk memicu sensor lazy-load
                            await page.keyboard.press("PageDown")
                            await page.keyboard.press("PageDown")
                            
                        except Exception:
                            pass

                    # Tunggu Gmaps meload data baru hasil scroll agresif tadi
                    if random.random() < 0.75:
                        await asyncio.sleep(random.uniform(1.0, 1.8))
                    else:
                        await asyncio.sleep(random.uniform(1.81, 3.0))

                    if extracted_in_this_cycle == 0:
                        retries += 1
                        if retries > 10: # Jika 10 kali percobaan scroll agresif tidak membuahkan data baru, berarti data habis.
                            break
                    else:
                        retries = 0

        except Exception as e: 
            console.print(f"[bold red]Error pada Review Scraper: {e}[/bold red]")
        finally:
            await context.close()

    def save_data(self, formats: List[str]) -> List[str]:
        saved_files = []
        if not self.results: 
            console.print("[bold yellow]Tidak ada data yang dapat diekspor![/bold yellow]")
            return saved_files
        
        # --- TAMBAHAN LOGIKA FALLBACK FORMAT ---
        if not formats or len(formats) == 0:
            formats = ["CSV"]
            console.print("[bold yellow]⚠️ Anda tidak mencentang format apapun. Sistem otomatis menyimpan sebagai .CSV[/bold yellow]")
            
        # Ganti '&' dengan 'dan' untuk menghindari bug visual Telegram Desktop
        safe_place_name = self.place_name.replace('&', 'dan')
        safe_place_name = re.sub(r'[\\/*?:"<>|]', "", safe_place_name).strip()
        if not safe_place_name:
            safe_place_name = "Google_Maps"
        
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        # Format baru: Nama Tempat Review_timestamp
        base_filename = os.path.join(self.output_dir, f"{safe_place_name}_reviews_{timestamp}")
        abs_dir = os.path.abspath(self.output_dir)
        # --- PERUBAHAN SELESAI DI SINI ---
        
        console.print(f"\n[bold green]✓ Selesai mengekstrak {len(self.results)} Ulasan unik![/bold green]")
        console.print(f"[bold cyan]📁 File Anda disimpan di folder: {abs_dir}[/bold cyan]\n")
        
        if "JSON" in formats:
            path = f"{base_filename}.json"
            with open(path, 'w', encoding='utf-8') as f: json.dump(self.results, f, ensure_ascii=False, indent=4)
            console.print(f"📄 JSON  : {path}")
            saved_files.append(path)
            
        if "CSV" in formats:
            path = f"{base_filename}.csv"
            with open(path, 'w', newline='', encoding='utf-8') as f:
                w = csv.DictWriter(f, fieldnames=self.headers)
                w.writeheader()
                w.writerows(self.results)
            console.print(f"📊 CSV   : {path}")
            saved_files.append(path)
            
        if "Excel (XLSX)" in formats:
            path = f"{base_filename}.xlsx"
            wb = Workbook()
            ws = wb.active
            ws.append(self.headers)
            for cell in ws[1]: cell.font = Font(bold=True)
            for item in self.results: ws.append([item.get(h, "N/A") for h in self.headers])

            for column_cells in ws.columns:
                max_length = 0
                column_letter = get_column_letter(column_cells[0].column)
                for cell in column_cells:
                    try:
                        if cell.value:
                            max_length = max(max_length, len(str(cell.value)))
                    except:
                        pass
                adjusted_width = (max_length + 2)
                ws.column_dimensions[column_letter].width = adjusted_width

            wb.save(path)
            console.print(f"📗 Excel : {path}")
            saved_files.append(path)
            
        if "Word (DOCX)" in formats:
            path = f"{base_filename}.docx"
            doc = docx.Document()
            doc.add_heading('Data Ulasan Google Maps', 0)
            for item in self.results:
                doc.add_heading(item.get("Reviewer Name", "Anonim"), level=2)
                for key in self.headers:
                    if key == "Reviewer Name": continue
                    p = doc.add_paragraph()
                    p.add_run(f"{key}: ").bold = True
                    p.add_run(str(item.get(key, "N/A")))
            doc.save(path)
            console.print(f"📘 Word  : {path}")
            saved_files.append(path)
        print("\n")
        return saved_files