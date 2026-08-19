# scraper.py
import asyncio
import csv
import json
import os
import random
import re
import sys
import urllib.parse
from datetime import datetime
from typing import List, Dict, Any, Optional, Callable

import aiohttp
import questionary
from bs4 import BeautifulSoup
from playwright.async_api import async_playwright, Page, BrowserContext, Browser, TimeoutError as PlaywrightTimeoutError
from rich.console import Console
from rich.panel import Panel
from rich.text import Text
from rich.progress import Progress, SpinnerColumn, TextColumn, BarColumn, TaskProgressColumn

import docx
from openpyxl import Workbook
from openpyxl.styles import Font
from openpyxl.utils import get_column_letter

console = Console()

USER_AGENTS = [
    "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36",
    "Mozilla/5.0 (Macintosh; Intel Mac OS X 10_15_7) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/119.0.0.0 Safari/537.36",
    "Mozilla/5.0 (Windows NT 10.0; Win64; x64; rv:121.0) Gecko/20100101 Firefox/121.0"
]

def add_hyperlink(paragraph, url, text):
    part = paragraph.part
    r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)
    hyperlink = docx.oxml.shared.OxmlElement('w:hyperlink')
    hyperlink.set(docx.oxml.shared.qn('r:id'), r_id, )
    new_run = docx.oxml.shared.OxmlElement('w:r')
    rPr = docx.oxml.shared.OxmlElement('w:rPr')
    c = docx.oxml.shared.OxmlElement('w:color')
    c.set(docx.oxml.shared.qn('w:val'), '0563C1')
    u = docx.oxml.shared.OxmlElement('w:u')
    u.set(docx.oxml.shared.qn('w:val'), 'single')
    rPr.append(c)
    rPr.append(u)
    new_run.append(rPr)
    new_run_text = docx.oxml.shared.OxmlElement('w:t')
    new_run_text.text = text
    new_run.append(new_run_text)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)
    return hyperlink

class GoogleMapsScraperAsync:
    def __init__(self, query: str, limit: int, headless: bool, proxy: Optional[str] = None, output_dir: str = "output", browser: Optional[Browser] = None, progress_callback: Optional[Callable[[int, int, str], Any]] = None) -> None:
        self.query = query
        self.limit = limit
        self.headless = headless
        self.proxy = proxy
        self.proxy_server = {"server": proxy} if proxy else None
        self.browser = browser
        self.results: List[Dict[str, Any]] = []
        self.output_dir = output_dir
        self.checkpoint_file = os.path.join(self.output_dir, "checkpoint_sementara.csv")
        self.semaphore = asyncio.Semaphore(5) 
        self.progress_callback = progress_callback
        
        if not os.path.exists(self.output_dir):
            os.makedirs(self.output_dir)

        self.headers = [
            "Name", "Rating", "Reviews", "Category", 
            "Google Maps", "Address", "Status/Hours", 
            "Latitude", "Longitude", "Phone", "Website", 
            "Instagram", "Facebook", "LinkedIn"
        ]
        
        with open(self.checkpoint_file, 'w', newline='', encoding='utf-8') as f:
            writer = csv.DictWriter(f, fieldnames=self.headers)
            writer.writeheader()

    async def _report_progress(self, current: int, total: int, phase_text: str = "") -> None:
        if self.progress_callback:
            try:
                if asyncio.iscoroutinefunction(self.progress_callback):
                    await self.progress_callback(current, total, phase_text)
                else:
                    self.progress_callback(current, total, phase_text)
            except Exception:
                pass

    async def _human_delay(self, min_sec: float = 1.0, max_sec: float = 2.5) -> None:
        await asyncio.sleep(random.uniform(min_sec, max_sec))

    async def _extract_socials(self, url: str) -> Dict[str, str]:
        socials = {"Instagram": "N/A", "Facebook": "N/A", "LinkedIn": "N/A"}
        if not url or url == "N/A": return socials
        try:
            async with aiohttp.ClientSession() as session:
                async with session.get(url, timeout=10) as response:
                    html = await response.text()
                    soup = BeautifulSoup(html, 'html.parser')
                    links = [a.get('href', '') for a in soup.find_all('a', href=True)]
                    for link in links:
                        link_lower = link.lower()
                        if 'instagram.com' in link_lower and socials["Instagram"] == "N/A": socials["Instagram"] = link
                        if 'facebook.com' in link_lower and socials["Facebook"] == "N/A": socials["Facebook"] = link
                        if 'linkedin.com' in link_lower and socials["LinkedIn"] == "N/A": socials["LinkedIn"] = link
        except: pass
        return socials

    def _save_checkpoint(self, data: Dict[str, str]):
        with open(self.checkpoint_file, 'a', newline='', encoding='utf-8') as f:
            writer = csv.DictWriter(f, fieldnames=self.headers)
            writer.writerow(data)

    async def _extract_details(self, context: BrowserContext, listing_url: str, progress_task, progress_bar, total_urls: int) -> None:
        async with self.semaphore:
            page = await context.new_page()
            data = {k: "N/A" for k in self.headers}
            data["Google Maps"] = listing_url

            try:
                await page.goto(listing_url, wait_until="domcontentloaded", timeout=30000)
                
                try:
                    await page.locator('h1').first.wait_for(state="visible", timeout=8000)
                    await self._human_delay(1.0, 2.0)
                except: pass

                # Lat & Long
                try:
                    match_3d = re.search(r'!3d(-?\d+\.\d+)!4d(-?\d+\.\d+)', listing_url)
                    if match_3d:
                        data["Latitude"] = match_3d.group(1)
                        data["Longitude"] = match_3d.group(2)
                    else:
                        match_at = re.search(r'@(-?\d+\.\d+),(-?\d+\.\d+)', page.url)
                        if match_at:
                            data["Latitude"] = match_at.group(1)
                            data["Longitude"] = match_at.group(2)
                except: pass

                # --- JAVASCRIPT INJECTION: Ekstraksi Data Secara Bersamaan (10x Lebih Cepat) ---
                try:
                    js_code = """
                    () => {
                        const h1 = document.querySelector('h1');
                        const h1_parent = h1?.parentElement?.parentElement?.innerText || "";
                        const address_btn = document.querySelector('button[data-item-id="address"]');
                        const address = address_btn ? address_btn.innerText : "";
                        
                        const status_divs = Array.from(document.querySelectorAll('div[class*="fontBodyMedium"]'));
                        const status = status_divs.find(e => e.innerText.includes("Buka") || e.innerText.includes("Tutup"))?.innerText || "";
                        
                        const phone_btn = document.querySelector('button[data-item-id^="phone:tel:"]');
                        const phone = phone_btn ? phone_btn.innerText : "";
                        
                        const web_a = document.querySelector('a[data-item-id="authority"]');
                        const website = web_a ? web_a.getAttribute('href') : "";
                        
                        const merchant_a = document.querySelector('a[data-item-id="merchant"]');
                        const merchant = merchant_a ? merchant_a.innerText : "";
                        
                        return {
                            name: h1 ? h1.innerText : "",
                            h1_parent, address, status, phone, website, merchant
                        };
                    }
                    """
                    extracted = await page.evaluate(js_code)
                except:
                    extracted = {"name":"", "h1_parent":"", "address":"", "status":"", "phone":"", "website":"", "merchant":""}

                # Nama
                if extracted["name"]:
                    data["Name"] = extracted["name"].strip()

                # FIX UTAMA: Rating, Reviews, dan Category menggunakan Smart Parsing
                try:
                    if extracted["h1_parent"]:
                        lines = [line.strip() for line in extracted["h1_parent"].split('\n') if line.strip()]
                        flat_text = " ".join(lines)
                    
                    # 1. Ambil Rating
                    match_rating = re.search(r'([1-5][.,][0-9])', flat_text)
                    if match_rating: data["Rating"] = match_rating.group(1).replace(',', '.')
                    
                    # 2. Ambil Ulasan
                    match_reviews = re.search(r'\(([\d.,]+)\)', flat_text)
                    if match_reviews: data["Reviews"] = match_reviews.group(1).replace('.', '').replace(',', '')

                    # 3. Ambil Kategori (Ambil baris tepat setelah baris yang mengandung ulasan)
                    for i, line in enumerate(lines):
                        if re.search(r'\(([\d.,]+)\)', line):
                            if i + 1 < len(lines):
                                data["Category"] = lines[i+1].split('·')[0].strip()
                            break
                    
                    # Fallback jika tidak ada ulasan
                    if data["Category"] == "N/A":
                        for i, line in enumerate(lines):
                            if re.search(r'([1-5][.,][0-9])', line):
                                if i + 1 < len(lines):
                                    data["Category"] = lines[i+1].split('·')[0].strip()
                                break
                except: pass

                # Alamat
                try: 
                    if extracted["address"]:
                        clean_address = extracted["address"].replace('', '').replace('?', '')
                        clean_address = ' '.join([x.strip() for x in clean_address.split('\n') if x.strip()])
                        data["Address"] = clean_address
                except: pass

                # Status/Jam Operasional
                try:
                    if extracted["status"]:
                        clean_status = extracted["status"].replace('', '').replace('keyboard_arrow_down', '')
                        clean_status = ' '.join([x.strip() for x in clean_status.split('\n') if x.strip()])
                        data["Status/Hours"] = clean_status
                except: pass

                # Telepon
                try: 
                    if extracted["phone"]:
                        clean_phone = extracted["phone"].replace('', '').replace('?', '')
                        clean_phone = ' '.join([x.strip() for x in clean_phone.split('\n') if x.strip()])
                        clean_phone = re.sub(r'[^\d\+\-\s\(\)]', '', clean_phone).strip()
                        data["Phone"] = clean_phone
                except: pass

                # Website & Sosmed
                try: 
                    if extracted["website"]:
                        data["Website"] = extracted["website"]
                        
                    # Ekstraksi Sosmed dari locator HTML
                    social_links = await page.locator('a[data-item-id="authority"]').evaluate_all('elements => elements.map(e => e.href)')
                    socials = {"Instagram": "N/A", "Facebook": "N/A", "LinkedIn": "N/A"}
                    for link in social_links:
                        link_lower = link.lower()
                        if 'instagram.com' in link_lower: socials["Instagram"] = link
                        elif 'facebook.com' in link_lower: socials["Facebook"] = link
                        elif 'linkedin.com' in link_lower: socials["LinkedIn"] = link
                    data.update(socials)
                except: pass
                
                self.results.append(data)
                self._save_checkpoint(data)
                
            except Exception as e:
                pass
            finally:
                await page.close()
                progress_bar.advance(progress_task)
                await self._report_progress(len(self.results), total_urls, f"Mengekstrak detail tempat ({len(self.results)}/{total_urls})...")

    async def scrape(self) -> None:
        if self.browser:
            context = await self.browser.new_context(
                viewport={"width": 1280, "height": 800},
                user_agent=random.choice(USER_AGENTS)
            )
            await self._run_scrape(context)
        else:
            async with async_playwright() as p:
                browser_args = ['--disable-blink-features=AutomationControlled']
                browser: Browser = await p.chromium.launch(headless=self.headless, args=browser_args, proxy=self.proxy_server)
                context: BrowserContext = await browser.new_context(
                    viewport={"width": 1280, "height": 800},
                    user_agent=random.choice(USER_AGENTS)
                )
                await self._run_scrape(context)
                await browser.close()
                
    async def _run_scrape(self, context: BrowserContext) -> None:
        async def _block_heavy_resources(route):
            if route.request.resource_type in ["image", "font", "media"]:
                await route.abort()
            else:
                await route.continue_()

        try:
            await context.route("**/*", _block_heavy_resources)
        except Exception:
            pass

        page: Page = await context.new_page()

        console.print("[bold green]🔍 Membuka Google Maps & Melakukan Pencarian (Scrolling)...[/bold green]")
        try:
            encoded_query = urllib.parse.quote_plus(self.query)
            search_url = f"https://www.google.com/maps/search/{encoded_query}"
            
            await page.goto(search_url, wait_until="domcontentloaded", timeout=60000)
            await self._human_delay(2.0, 3.0)

            feed_locator = page.locator('div[role="feed"]')
            await feed_locator.wait_for(state="attached", timeout=15000)
            
            prev_count = 0
            retries = 0
            while True:
                elements = page.locator('a[href*="https://www.google.com/maps/place/"]')
                current_count = await elements.count()
                
                console.print(f"[cyan]Menemukan {current_count} lokasi (Target: {self.limit})...[/cyan]", end="\r")
                await self._report_progress(min(current_count, self.limit), self.limit, f"Menemukan {current_count}/{self.limit} lokasi...")
                if current_count >= self.limit: break
                
                if current_count == prev_count:
                    retries += 1
                    if retries > 10:
                        break
                else:
                    retries = 0
                    
                prev_count = current_count
                await feed_locator.evaluate("element => element.scrollTop = element.scrollHeight")
                if random.random() < 0.75:
                    await asyncio.sleep(random.uniform(1.0, 1.8))
                else:
                    await asyncio.sleep(random.uniform(1.81, 3.0))
            
            console.print("\n")
            all_links = await page.locator('a[href*="https://www.google.com/maps/place/"]').all()
            listing_urls = []
            for link in all_links[:self.limit]:
                href = await link.get_attribute("href")
                if href: listing_urls.append(href)
            
            await page.close() 

            console.print(f"[bold magenta]⚡ Memulai Ekstraksi Asinkron ({len(listing_urls)} data)...[/bold magenta]")
            await self._report_progress(0, len(listing_urls), f"Mulai mengekstrak {len(listing_urls)} detail tempat...")
            
            is_bot = os.environ.get('IS_TELEGRAM_BOT') == '1'
            if is_bot:
                class DummyProgress:
                    def add_task(self, *args, **kwargs): return None
                    def advance(self, *args, **kwargs): pass
                    def __enter__(self): return self
                    def __exit__(self, *args): pass
                progress_ctx = DummyProgress()
            else:
                progress_ctx = Progress(
                    SpinnerColumn(),
                    TextColumn("[progress.description]{task.description}"),
                    BarColumn(),
                    TaskProgressColumn(),
                )
                
            with progress_ctx as progress:
                extract_task = progress.add_task("[cyan]Mengekstrak Detail & Sosmed...", total=len(listing_urls))
                
                tasks = [self._extract_details(context, url, extract_task, progress, len(listing_urls)) for url in listing_urls]
                await asyncio.gather(*tasks)

        except Exception as e:
            console.print(f"[bold red]Scraping Error: {e}[/bold red]")


    def save_data(self, formats: List[str]) -> List[str]:
        saved_files = []
        if not self.results: return saved_files

        # --- TAMBAHAN LOGIKA FALLBACK FORMAT ---
        if not formats or len(formats) == 0:
            formats = ["CSV"]
            console.print("[bold yellow]⚠️ Anda tidak mencentang format apapun. Sistem otomatis menyimpan sebagai .CSV[/bold yellow]")
        
        # Ganti '&' dengan 'dan' untuk menghindari bug visual Telegram Desktop
        safe_query = self.query.replace('&', 'dan')
        safe_query = re.sub(r'[\\/*?:"<>|]', "", safe_query).strip()
        
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        # Format baru: Query_results_timestamp
        base_filename = os.path.join(self.output_dir, f"{safe_query}_results_{timestamp}")
        # --- PERUBAHAN SELESAI DI SINI ---
        
        console.print(f"\n[bold green]✓ Selesai mengekstrak {len(self.results)} Tempat![/bold green]")
        
        if "JSON" in formats:
            path = f"{base_filename}.json"
            with open(path, 'w', encoding='utf-8') as f: json.dump(self.results, f, ensure_ascii=False, indent=4)
            console.print(f"📄 JSON  : [cyan]{path}[/cyan]")
            saved_files.append(path)

        if "CSV" in formats:
            path = f"{base_filename}.csv"
            with open(path, 'w', newline='', encoding='utf-8') as f:
                writer = csv.DictWriter(f, fieldnames=self.headers)
                writer.writeheader()
                writer.writerows(self.results)
            console.print(f"📊 CSV   : [cyan]{path}[/cyan]")
            saved_files.append(path)

        if "Excel (XLSX)" in formats:
            path = f"{base_filename}.xlsx"
            wb = Workbook()
            ws = wb.active
            ws.append(self.headers)
            for cell in ws[1]: cell.font = Font(bold=True)
                
            for item in self.results:
                row_data = [item.get(h, "N/A") for h in self.headers]
                ws.append(row_data)
                current_row = ws.max_row
                
                link_columns = {
                    "Google Maps": "Buka Maps",
                    "Website": "Buka Website",
                    "Instagram": "Buka Instagram",
                    "Facebook": "Buka Facebook",
                    "LinkedIn": "Buka LinkedIn"
                }
                
                for key_link, text_display in link_columns.items():
                    if key_link in self.headers:
                        idx = self.headers.index(key_link) + 1
                        cell = ws.cell(row=current_row, column=idx)
                        # Biarkan "N/A" tetap N/A, hanya jadikan hyperlink jika ada URL
                        if cell.value != "N/A" and cell.value:
                            url = cell.value
                            cell.value = text_display
                            cell.hyperlink = url
                            cell.font = Font(color="0563C1", underline="single")

            for column_cells in ws.columns:
                max_length = 0
                column_letter = get_column_letter(column_cells[0].column)
                for cell in column_cells:
                    try:
                        if cell.value:
                            max_length = max(max_length, len(str(cell.value)))
                    except:
                        pass
                adjusted_width = (max_length + 2)
                ws.column_dimensions[column_letter].width = adjusted_width

            wb.save(path)
            console.print(f"📗 Excel : [cyan]{path}[/cyan]")
            saved_files.append(path)

        if "Word (DOCX)" in formats:
            path = f"{base_filename}.docx"
            doc = docx.Document()
            doc.add_heading(f'Data Enrichment: {self.query}', 0)
            
            link_columns = ["Google Maps", "Website", "Instagram", "Facebook", "LinkedIn"]
            
            for item in self.results:
                doc.add_heading(item.get("Name", "N/A"), level=1)
                
                for key in self.headers:
                    if key == "Name": continue
                    
                    p = doc.add_paragraph()
                    p.add_run(f"{key}: ").bold = True
                    
                    if key in link_columns:
                        if item.get(key) and item.get(key) != "N/A":
                            add_hyperlink(p, item[key], f"Buka {key.split()[-1] if 'Link' in key else key}")
                        else:
                            p.add_run("N/A")
                    else:
                        p.add_run(str(item.get(key, "N/A")))
                        
            doc.save(path)
            console.print(f"📘 Word  : [cyan]{path}[/cyan]")
            saved_files.append(path)
        print("\n")
        return saved_files